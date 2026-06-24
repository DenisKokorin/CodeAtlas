import io
import json
import os
import re
from datetime import datetime, timezone
from typing import Literal

from app.models import Repository

ExportFormat = Literal["markdown", "txt", "html", "docx", "json"]

_BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_TEMPLATE_DIR = os.path.join(_BASE_DIR, "templates")
_TEMPLATE_PATH = os.path.join(_TEMPLATE_DIR, "export.html")


def _load_html_template() -> str:
    try:
        with open(_TEMPLATE_PATH, encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        raise RuntimeError(
            f"Export HTML template not found at {_TEMPLATE_PATH}"
        )


def _build_meta_html(repository: Repository) -> str:
    pieces = []

    if repository.github_full_name:
        pieces.append(
            f'<span class="meta-item">'
            f'<span class="meta-label">Repository:</span> '
            f'<a href="{repository.repo_url}">{repository.github_full_name}</a>'
            f"</span>"
        )

    if repository.github_language:
        pieces.append(
            f'<span class="meta-item">'
            f'<span class="meta-label">Language:</span> {repository.github_language}'
            f"</span>"
        )

    if repository.github_default_branch:
        pieces.append(
            f'<span class="meta-item">'
            f'<span class="meta-label">Branch:</span> '
            f"{repository.github_default_branch}"
            f"</span>"
        )

    if repository.documentation_provider:
        pieces.append(
            f'<span class="meta-item">'
            f'<span class="meta-label">Provider:</span> '
            f"{repository.documentation_provider}"
            f"</span>"
        )

    if repository.documentation_updated_at:
        updated_str = repository.documentation_updated_at.strftime(
            "%Y-%m-%d %H:%M UTC"
        )
        pieces.append(
            f'<span class="meta-item">'
            f'<span class="meta-label">Updated:</span> {updated_str}'
            f"</span>"
        )

    if not pieces:
        return ""

    return f'<div class="meta">{"".join(pieces)}</div>'


def _build_meta_text(repository: Repository) -> str:
    lines: list[str] = []

    if repository.github_full_name:
        lines.append(f"Repository: {repository.github_full_name}")
    if repository.github_language:
        lines.append(f"Language: {repository.github_language}")
    if repository.documentation_provider:
        lines.append(f"Provider: {repository.documentation_provider}")
    if repository.documentation_updated_at:
        lines.append(
            f"Updated: {repository.documentation_updated_at.strftime('%Y-%m-%d %H:%M UTC')}"
        )

    return "\n".join(lines) if lines else ""


def _markdown_to_plain_text(md: str) -> str:
    text = md

    text = re.sub(
        r"```\w*\n([\s\S]*?)```",
        lambda m: m.group(1).strip(),
        text,
    )

    text = re.sub(r"`([^`]+)`", r"\1", text)

    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)

    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    
    text = re.sub(r"([*_]{1,3})([^*_]+?)\1", r"\2", text)
    
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    
    text = re.sub(r"^[\t ]*[-*+]\s+", "  • ", text, flags=re.MULTILINE)
    
    text = re.sub(r"^[\t ]*\d+\.\s+", "  ", text, flags=re.MULTILINE)
    
    text = re.sub(r"^[|\s:\-]+$", "", text, flags=re.MULTILINE)
    
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    return text.strip()


def _convert_markdown_to_html(markdown_text: str) -> str:
    import markdown as md_lib

    extensions = [
        "extra",
        "codehilite",
        "toc",
        "sane_lists",
    ]
    return md_lib.markdown(
        markdown_text,
        extensions=extensions,
        output_format="html5",
    )


def export_as_markdown(
    repository: Repository,
) -> tuple[str, str, str]:

    content = repository.generated_documentation or ""
    filename = f"{repository.name}-documentation.md"
    return content, "text/markdown; charset=utf-8", filename


def export_as_html(
    repository: Repository,
) -> tuple[str, str, str]:

    raw_md = repository.generated_documentation or ""
    html_body = _convert_markdown_to_html(raw_md)
    meta_block = _build_meta_html(repository)
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    template = _load_html_template()

    page = template
    page = page.replace("{{ title }}", repository.name)
    page = page.replace("{{ meta_html }}", meta_block)
    page = page.replace("{{ content }}", html_body)
    page = page.replace("{{ date }}", now_str)

    filename = f"{repository.name}-documentation.html"
    return page, "text/html; charset=utf-8", filename


def export_as_json(
    repository: Repository,
) -> tuple[str, str, str]:

    payload = {
        "repository": {
            "id": repository.id,
            "name": repository.name,
            "url": repository.repo_url,
            "full_name": repository.github_full_name,
            "description": repository.description,
            "language": repository.github_language,
            "default_branch": repository.github_default_branch,
        },
        "documentation": {
            "content": repository.generated_documentation,
            "provider": repository.documentation_provider,
            "is_stale": repository.documentation_is_stale,
            "updated_at": (
                repository.documentation_updated_at.isoformat()
                if repository.documentation_updated_at
                else None
            ),
            "source_updated_at": repository.documentation_source_updated_at,
        },
        "export": {
            "format": "json",
            "generated_at": datetime.now(timezone.utc).isoformat(),
        },
    }

    content = json.dumps(payload, indent=2, ensure_ascii=False)
    filename = f"{repository.name}-documentation.json"
    return content, "application/json; charset=utf-8", filename


def export_as_txt(
    repository: Repository,
) -> tuple[str, str, str]:

    raw_md = repository.generated_documentation or ""
    body = _markdown_to_plain_text(raw_md)
    meta = _build_meta_text(repository)

    header = f"# {repository.name}\n"
    if meta:
        header += f"\n{meta}\n"
    header += "\n"

    content = header + body + "\n"
    filename = f"{repository.name}-documentation.txt"
    return content, "text/plain; charset=utf-8", filename


def export_as_docx(
    repository: Repository,
) -> tuple[bytes, str, str]:

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_heading(repository.name, level=1)

    if repository.github_full_name:
        p = doc.add_paragraph()
        run = p.add_run("Repository: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        run = p.add_run(repository.github_full_name)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    if repository.github_language:
        p = doc.add_paragraph()
        run = p.add_run("Language: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        run = p.add_run(repository.github_language)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    if repository.documentation_provider:
        p = doc.add_paragraph()
        run = p.add_run("Provider: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        run = p.add_run(repository.documentation_provider)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    if repository.documentation_updated_at:
        p = doc.add_paragraph()
        run = p.add_run("Updated: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        run = p.add_run(
            repository.documentation_updated_at.strftime("%Y-%m-%d %H:%M UTC")
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("_" * 60)

    raw_md = repository.generated_documentation or ""
    _feed_markdown_to_docx(doc, raw_md)

    doc.add_paragraph("")  # spacer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    run = p.add_run(f"Generated by CodeAtlas - {now_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(156, 163, 175)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{repository.name}-documentation.docx"
    return buf.getvalue(), (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ), filename


def _feed_markdown_to_docx(doc, md: str) -> None:

    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    lines = md.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_data: list[list[str]] = []
    list_type: str | None = None

    def flush_code():
        nonlocal code_lines, in_code_block
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            text = "\n".join(code_lines)
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 41, 59)
            shd = p.paragraph_format.element.makeelement(
                qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F1F5F9"}
            )
            pPr = p.paragraph_format.element.get_or_add_pPr()
            pPr.append(shd)
            code_lines = []

    def flush_table():
        nonlocal table_data, in_table
        if not table_data:
            return
        rows = len(table_data)
        if rows == 0:
            return
        cols = max(len(r) for r in table_data)
        if cols == 0:
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Light Grid Accent 1"
        for r_idx, row_data in enumerate(table_data):
            for c_idx in range(cols):
                cell = table.cell(r_idx, c_idx)
                cell.text = row_data[c_idx] if c_idx < len(row_data) else ""

        table_data = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```") and not in_code_block:
            in_code_block = True
            i += 1
            continue
        if line.startswith("```") and in_code_block:
            in_code_block = False
            flush_code()
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            if re.match(r"^[\s|:\-]+$", line):
                i += 1
                continue
            table_data.append(
                [c.strip() for c in line.strip().strip("|").split("|")]
            )
            in_table = True
            i += 1
            continue
        if in_table:
            flush_table()
            in_table = False

        h_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            if level <= 2:
                doc.add_heading(text, level=level)
            elif level == 3:
                h = doc.add_heading(text, level=3)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        if re.match(r"^[-*_]{3,}\s*$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p.paragraph_format.element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "CCCCCC"},
            )
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        bq_match = re.match(r"^>\s?(.*)$", line)
        if bq_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(bq_match.group(1))
            run.font.color.rgb = RGBColor(71, 85, 105)
            run.italic = True
            i += 1
            continue

        ul_match = re.match(r"^[\t ]*[-*+]\s+(.+)$", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            p.text = ul_match.group(1)
            i += 1
            continue

        ol_match = re.match(r"^[\t ]*\d+\.\s+(.+)$", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            p.text = ol_match.group(1)
            i += 1
            continue

        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            _add_inline_text(p, stripped)

        i += 1

    flush_code()
    if in_table:
        flush_table()


def _add_inline_text(paragraph, text: str) -> None:

    parts = re.split(r"(\*\*.+?\*\*|__.+?__)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r"(\*.+?\*|_.+?_)", part)
            for sub in sub_parts:
                if (sub.startswith("*") and sub.endswith("*")) or (
                    sub.startswith("_") and sub.endswith("_")
                ):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    code_parts = re.split(r"(`.+?`)", sub)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            run = paragraph.add_run(cp[1:-1])
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)
                        else:
                            paragraph.add_run(cp)


EXPORTERS: dict[ExportFormat, callable] = {
    "markdown": export_as_markdown,
    "txt": export_as_txt,
    "html": export_as_html,
    "docx": export_as_docx,
    "json": export_as_json,
}


def export_documentation(
    repository: Repository,
    fmt: ExportFormat,
) -> tuple[str | bytes, str, str]:

    exporter = EXPORTERS.get(fmt)

    if exporter is None:
        raise ValueError(
            f"Unsupported export format: {fmt!r}. "
            f"Available formats: {', '.join(EXPORTERS)}"
        )

    return exporter(repository)
